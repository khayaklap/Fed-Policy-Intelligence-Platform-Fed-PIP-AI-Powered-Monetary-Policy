"""
Format Exporters

Export reports to multiple formats: PDF, DOCX, HTML, Markdown.
"""

import logging
from typing import Dict, Optional
from datetime import datetime
import json

# Document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX export disabled")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("reportlab not available - PDF export disabled")

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logging.warning("markdown not available - some HTML features may be limited")

try:
    from jinja2 import Template
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False
    logging.warning("jinja2 not available - HTML export disabled")

try:
    # Try relative imports first (when used as module)
    from .report_generator_config import (
        DOCUMENT_STYLE,
        HEADER_FOOTER,
        PDF_SETTINGS,
        DOCX_SETTINGS,
        HTML_SETTINGS
    )
except ImportError:
    # Fall back to absolute imports (when run directly)
    from report_generator_config import (
        DOCUMENT_STYLE,
        HEADER_FOOTER,
        PDF_SETTINGS,
        DOCX_SETTINGS,
        HTML_SETTINGS
    )

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ReportExporter:
    """
    Export reports to various formats.
    """
    
    def __init__(self):
        """Initialize exporter."""
        logger.info("Initialized Report Exporter")
    
    def export_to_docx(self, report: Dict, filename: str) -> str:
        """
        Export report to DOCX format.
        
        Args:
            report: Report dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to generated file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        logger.info(f"Exporting to DOCX: {filename}")
        
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = report['metadata']['title']
        doc.core_properties.author = report['metadata']['author']
        doc.core_properties.created = datetime.fromisoformat(report['metadata']['generated_at'])
        
        # Add title
        title = doc.add_heading(report['metadata']['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Generated: {report['metadata']['generated_at']}\n").italic = True
        metadata_para.add_run(f"Report Type: {report['metadata']['report_type']}\n").italic = True
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Add each section
        for section in report['sections']:
            # Section heading
            doc.add_heading(section['title'], level=1)
            
            # Section content
            content = section.get('content', {})
            
            if section['type'] == 'summary':
                # Bullet points for summary
                for point in content.get('key_points', []):
                    doc.add_paragraph(point, style='List Bullet')
            
            elif section['type'] == 'analysis':
                # Formatted analysis
                for key, value in content.items():
                    if isinstance(value, (str, int, float)):
                        p = doc.add_paragraph()
                        p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                        p.add_run(str(value))
            
            elif section['type'] == 'data':
                # Table for data
                if content:
                    self._add_data_table(doc, content)
            
            elif section['type'] == 'recommendations':
                # Numbered list for recommendations
                for i, rec in enumerate(content.get('recommendations', []), 1):
                    doc.add_paragraph(f"{i}. {rec}")
            
            doc.add_paragraph()  # Spacing
        
        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"{report['metadata']['author']} | Page "
        
        # Save
        output_path = f"{filename}.docx"
        doc.save(output_path)
        
        logger.info(f"DOCX exported: {output_path}")
        return output_path
    
    def export_to_pdf(self, report: Dict, filename: str) -> str:
        """
        Export report to PDF format.
        
        Args:
            report: Report dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to generated file
        """
        if not PDF_AVAILABLE:
            raise ImportError("reportlab not installed")
        
        logger.info(f"Exporting to PDF: {filename}")
        
        output_path = f"{filename}.pdf"
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12
        )
        
        # Title
        story.append(Paragraph(report['metadata']['title'], title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        metadata_text = f"Generated: {report['metadata']['generated_at']}<br/>"
        metadata_text += f"Report Type: {report['metadata']['report_type']}"
        story.append(Paragraph(metadata_text, styles['Italic']))
        story.append(PageBreak())
        
        # Sections
        for section in report['sections']:
            # Section heading
            story.append(Paragraph(section['title'], heading_style))
            story.append(Spacer(1, 0.1*inch))
            
            # Content
            content = section.get('content', {})
            
            if section['type'] == 'summary':
                for point in content.get('key_points', []):
                    story.append(Paragraph(f"• {point}", styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
            
            elif section['type'] == 'analysis':
                for key, value in content.items():
                    text = f"<b>{key.replace('_', ' ').title()}:</b> {value}"
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
            
            elif section['type'] == 'recommendations':
                for i, rec in enumerate(content.get('recommendations', []), 1):
                    story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
            
            story.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"PDF exported: {output_path}")
        return output_path
    
    def export_to_html(self, report: Dict, filename: str) -> str:
        """
        Export report to HTML format.
        
        Args:
            report: Report dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to generated file
        """
        if not JINJA2_AVAILABLE:
            raise ImportError("jinja2 not installed")
        
        logger.info(f"Exporting to HTML: {filename}")
        
        # HTML template
        template = Template('''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }
        h1 {
            color: #1f77b4;
            border-bottom: 3px solid #1f77b4;
            padding-bottom: 10px;
        }
        h2 {
            color: #333;
            margin-top: 30px;
            border-bottom: 1px solid #ddd;
            padding-bottom: 5px;
        }
        .metadata {
            color: #666;
            font-style: italic;
            margin-bottom: 30px;
        }
        .summary-points {
            list-style-type: disc;
            padding-left: 20px;
        }
        .analysis-item {
            margin: 10px 0;
        }
        .key {
            font-weight: bold;
        }
        .recommendations {
            counter-reset: rec-counter;
            list-style: none;
            padding-left: 0;
        }
        .recommendations li {
            counter-increment: rec-counter;
            margin: 10px 0;
        }
        .recommendations li:before {
            content: counter(rec-counter) ". ";
            font-weight: bold;
            color: #1f77b4;
        }
        .footer {
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            color: #666;
            font-size: 0.9em;
        }
    </style>
</head>
<body>
    <h1>{{ title }}</h1>
    <div class="metadata">
        <p>Generated: {{ generated_at }}</p>
        <p>Report Type: {{ report_type }}</p>
    </div>
    
    {% for section in sections %}
    <h2>{{ section.title }}</h2>
    
    {% if section.type == 'summary' %}
        <ul class="summary-points">
        {% for point in section.content.key_points %}
            <li>{{ point }}</li>
        {% endfor %}
        </ul>
    
    {% elif section.type == 'analysis' %}
        {% for key, value in section.content.items() %}
        <div class="analysis-item">
            <span class="key">{{ key.replace('_', ' ').title() }}:</span> {{ value }}
        </div>
        {% endfor %}
    
    {% elif section.type == 'recommendations' %}
        <ol class="recommendations">
        {% for rec in section.content.recommendations %}
            <li>{{ rec }}</li>
        {% endfor %}
        </ol>
    
    {% endif %}
    {% endfor %}
    
    <div class="footer">
        <p>{{ author }} | {{ organization }}</p>
        <p>{{ disclaimer }}</p>
    </div>
</body>
</html>
        ''')
        
        # Render
        html_content = template.render(
            title=report['metadata']['title'],
            generated_at=report['metadata']['generated_at'],
            report_type=report['metadata']['report_type'],
            sections=report['sections'],
            author=report['metadata']['author'],
            organization=report['metadata'].get('organization', ''),
            disclaimer=report['metadata'].get('disclaimer', '')
        )
        
        # Save
        output_path = f"{filename}.html"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"HTML exported: {output_path}")
        return output_path
    
    def export_to_markdown(self, report: Dict, filename: str) -> str:
        """
        Export report to Markdown format.
        
        Args:
            report: Report dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to generated file
        """
        logger.info(f"Exporting to Markdown: {filename}")
        
        md_lines = []
        
        # Title
        md_lines.append(f"# {report['metadata']['title']}")
        md_lines.append("")
        
        # Metadata
        md_lines.append(f"**Generated:** {report['metadata']['generated_at']}  ")
        md_lines.append(f"**Report Type:** {report['metadata']['report_type']}  ")
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")
        
        # Sections
        for section in report['sections']:
            md_lines.append(f"## {section['title']}")
            md_lines.append("")
            
            content = section.get('content', {})
            
            if section['type'] == 'summary':
                for point in content.get('key_points', []):
                    md_lines.append(f"- {point}")
                md_lines.append("")
            
            elif section['type'] == 'analysis':
                for key, value in content.items():
                    md_lines.append(f"**{key.replace('_', ' ').title()}:** {value}  ")
                md_lines.append("")
            
            elif section['type'] == 'recommendations':
                for i, rec in enumerate(content.get('recommendations', []), 1):
                    md_lines.append(f"{i}. {rec}")
                md_lines.append("")
        
        # Footer
        md_lines.append("---")
        md_lines.append("")
        md_lines.append(f"*{report['metadata']['author']} | {report['metadata'].get('organization', '')}*")
        
        # Save
        output_path = f"{filename}.md"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_lines))
        
        logger.info(f"Markdown exported: {output_path}")
        return output_path
    
    def export_to_json(self, report: Dict, filename: str) -> str:
        """
        Export report to JSON format.
        
        Args:
            report: Report dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to generated file
        """
        logger.info(f"Exporting to JSON: {filename}")
        
        output_path = f"{filename}.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2)
        
        logger.info(f"JSON exported: {output_path}")
        return output_path
    
    def _add_data_table(self, doc: 'Document', data: Dict):
        """Add a data table to DOCX document."""
        if not data:
            return
        
        # Create table
        table = doc.add_table(rows=len(data)+1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        # Data
        for i, (key, value) in enumerate(data.items(), 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = key.replace('_', ' ').title()
            row_cells[1].text = str(value)


def export_report(report: Dict, filename: str, format: str = 'pdf') -> str:
    """
    Convenience function to export report.
    
    Args:
        report: Report dictionary
        filename: Output filename (without extension)
        format: Output format ('pdf', 'docx', 'html', 'markdown', 'json')
    
    Returns:
        Path to generated file
    """
    exporter = ReportExporter()
    
    if format == 'pdf':
        return exporter.export_to_pdf(report, filename)
    elif format == 'docx':
        return exporter.export_to_docx(report, filename)
    elif format == 'html':
        return exporter.export_to_html(report, filename)
    elif format == 'markdown' or format == 'md':
        return exporter.export_to_markdown(report, filename)
    elif format == 'json':
        return exporter.export_to_json(report, filename)
    else:
        raise ValueError(f"Unsupported format: {format}")
